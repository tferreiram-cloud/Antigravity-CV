
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown

# Configuration
INPUT_FILE = "master_profile_v8.json" # Now points to the mockup file
OUTPUT_DIR = "output_v83"
VERSION = "8.3-MOCKUP"

# MOCKUP Content
MULTIVERSE_VARIANTS = {
    "1. The AI Engineer": (
        "Experienced Engineer focused on automation. Architect of systems using Python and Docker. "
        "Built pipelines that reduced manual workload by 50%."
    ),
    "2. The Corporate Executive": (
        "Executive with 15+ years of experience in Fortune 500 companies. Specialist in "
        "corporate strategy and P&L management."
    ),
    "3. The Growth Hacker": (
        "Head of Growth obsessed with metrics. Scaling revenue and optimizing conversion funnels "
        "through A/B testing."
    )
}

TECH_STACK_DEEP_DIVE = {
    "AI Orchestration": {
        "n8n": [
            "Use Case: Orchestration of workflows.",
            "Complexity: Use of HTTP Request and Logic nodes."
        ],
        "Docker": [
            "Use Case: Containerization of apps.",
            "Ops: Management of volumes and networks."
        ]
    },
    "CRM & Automation": {
        "HubSpot": [
            "Architecture: Sales Pipelines and Lead Scoring."
        ],
        "Salesforce": [
            "Architecture: Dashboards and Reporting."
        ]
    }
}

ATOMIC_EXPERIENCES = {
    "TechCorp Solutions": [
        {
            "title": "Project A: Automation",
            "context": "Client needed efficiency.",
            "situation": "Manual processes were slow.",
            "task": "Automate everything.",
            "action_technical": "Implemented Python scripts.",
            "action_strategic": "Trained the team.",
            "result": "50% faster turnaround.",
            "keywords": ["Automation", "Python"]
        }
    ],
    "Global Systems Inc": [
        {
            "title": "Project A: Product Launch",
            "context": "New product needed.",
            "situation": "Market was ready.",
            "task": "Launch fast.",
            "action_technical": "Agile sprints.",
            "action_strategic": "GTM strategy.",
            "result": "Success.",
            "keywords": ["Product", "Agile"]
        }
    ]
}

ACADEMIC_APPLIED = {
    "University X": {
        "Focus": "Computer Science",
        "Application": "Applied algorithms to business problems."
    }
}

KEYWORD_DICTIONARY = {
    "Hard Skills": ["Python", "SQL", "Docker", "AWS"],
    "Soft Skills": ["Leadership", "Communication"],
    "Methodologies": ["Agile", "Scrum"]
}

def generate_v83():
    # 1. Hardcoded Base Data (Mockup)
    base_data = {
        "meta": {"version": "8.0-MOCKUP"},
        "candidato": {
            "nome_completo": "Alex Silva (Mockup)",
            "localizacao": {"cidade": "SP", "pais": "Brasil"},
            "contato": {"email": "alex@example.com"}
        }
    }

    # 2. Merge Data
    full_data = base_data.copy()
    full_data['multiverse_variants'] = MULTIVERSE_VARIANTS
    full_data['tech_stack_details'] = TECH_STACK_DEEP_DIVE
    full_data['atomic_experiences'] = ATOMIC_EXPERIENCES
    full_data['academic_applied'] = ACADEMIC_APPLIED
    full_data['keyword_dictionary'] = KEYWORD_DICTIONARY
    
    # 3. Create Output Directory
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    # 4. Save JSON
    json_path = os.path.join(OUTPUT_DIR, "master_cv_database_v8.3.json")
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(full_data, f, indent=4, ensure_ascii=False)
    print(f"JSON saved to {json_path}")

    # 5. Generate DOCX
    generate_docx(full_data)

    # 6. Generate PDF (via HTML)
    generate_pdf(full_data)

def generate_docx(data):
    doc = Document()
    
    # Title
    doc.add_heading('MASTER CV DATABASE (V8.3 - MOCKUP)', 0)
    doc.add_paragraph(f"Target: Public Repo | Version: {VERSION}")
    
    # Part 1: The Multiverse
    doc.add_heading('PART 1: THE MULTIVERSE', level=1)
    
    for title, text in MULTIVERSE_VARIANTS.items():
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    doc.add_page_break()
    
    # Part 2: Tech Stack Deep Dive
    doc.add_heading('PART 2: TECH STACK DEEP DIVE', level=1)
    
    for category, stacks in TECH_STACK_DEEP_DIVE.items():
        doc.add_heading(category, level=2)
        for tool, details in stacks.items():
            doc.add_heading(tool, level=3)
            for detail in details:
                doc.add_paragraph(detail, style='List Bullet')
                
    doc.add_page_break()
    
    # Part 3: Atomic Experience
    doc.add_heading('PART 3: ATOMIC EXPERIENCE', level=1)
    
    for company, projects in ATOMIC_EXPERIENCES.items():
        doc.add_heading(company.upper(), level=2)
        for proj in projects:
            doc.add_heading(proj['title'], level=3)
            doc.add_paragraph(f"Context: {proj['context']}")
            doc.add_paragraph(f"Situation: {proj['situation']}")
            doc.add_paragraph(f"Task: {proj['task']}")
            doc.add_paragraph(f"Action (Technical): {proj['action_technical']}")
            doc.add_paragraph(f"Action (Strategic): {proj['action_strategic']}")
            doc.add_paragraph(f"Result: {proj['result']}")
            doc.add_paragraph(f"Keywords: {', '.join(proj['keywords'])}")
            doc.add_paragraph("-" * 20)
            
    doc.add_page_break()
    
    # Part 4: Academic
    doc.add_heading('PART 4: ACADEMIC', level=1)
    for institution, details in ACADEMIC_APPLIED.items():
        doc.add_heading(institution, level=2)
        doc.add_paragraph(f"Focus: {details['Focus']}")
        doc.add_paragraph(f"Application: {details['Application']}")
        
    doc.add_page_break()
    
    # Part 5: Keywords
    doc.add_heading('PART 5: KEYWORD DICTIONARY', level=1)
    for category, words in KEYWORD_DICTIONARY.items():
        doc.add_heading(category, level=2)
        doc.add_paragraph(", ".join(words))

    docx_path = os.path.join(OUTPUT_DIR, "master_cv_database_v8.3.docx")
    doc.save(docx_path)
    print(f"DOCX saved to {docx_path}")

def generate_pdf(data):
    # Create HTML string
    html = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Helvetica, Arial, sans-serif; font-size: 10pt; line-height: 1.4; }}
            h1 {{ color: #2c3e50; border-bottom: 2px solid #2c3e50; padding-bottom: 10px; margin-top: 30px; }}
            h2 {{ color: #34495e; margin-top: 20px; }}
            h3 {{ color: #16a085; margin-top: 15px; }}
            .page-break {{ page-break-before: always; }}
            .meta {{ font-size: 8pt; color: #7f8c8d; text-align: center; margin-bottom: 30px; }}
            .project {{ background: #f9f9f9; padding: 15px; border-left: 4px solid #16a085; margin-bottom: 15px; }}
            .label {{ font-weight: bold; color: #2c3e50; }}
        </style>
    </head>
    <body>
        <div style="text-align: center;">
            <h1>MASTER CV DATABASE (V{VERSION} - MOCKUP)</h1>
            <p class="meta">Architecture: Mockup for Public Repo</p>
        </div>

        <h1>PART 1: THE MULTIVERSE</h1>
    """
    
    for title, text in MULTIVERSE_VARIANTS.items():
        html += f"<h3>{title}</h3><p>{text}</p>"
        
    html += "<div class='page-break'></div><h1>PART 2: TECH STACK DEEP DIVE</h1>"
    
    for category, stacks in TECH_STACK_DEEP_DIVE.items():
        html += f"<h2>{category}</h2>"
        for tool, details in stacks.items():
            html += f"<h3>{tool}</h3><ul>"
            for detail in details:
                html += f"<li>{detail}</li>"
            html += "</ul>"

    html += "<div class='page-break'></div><h1>PART 3: ATOMIC EXPERIENCE</h1>"
    
    for company, projects in ATOMIC_EXPERIENCES.items():
        html += f"<h2>{company}</h2>"
        for proj in projects:
            html += f"""
            <div class="project">
                <h3>{proj['title']}</h3>
                <p><span class="label">Context:</span> {proj['context']}</p>
                <p><span class="label">Situation:</span> {proj['situation']}</p>
                <p><span class="label">Task:</span> {proj['task']}</p>
                <p><span class="label">Action (Technical):</span> {proj['action_technical']}</p>
                <p><span class="label">Action (Strategic):</span> {proj['action_strategic']}</p>
                <p><span class="label">Result:</span> {proj['result']}</p>
                <p><span class="label">Keywords:</span> {', '.join(proj['keywords'])}</p>
            </div>
            """

    html += "<div class='page-break'></div><h1>PART 4: ACADEMIC</h1>"
    for institution, details in ACADEMIC_APPLIED.items():
        html += f"<h3>{institution}</h3>"
        html += f"<p><span class='label'>Focus:</span> {details['Focus']}</p>"
        html += f"<p><span class='label'>Application:</span> {details['Application']}</p>"
        
    html += "<div class='page-break'></div><h1>PART 5: KEYWORD DICTIONARY</h1>"
    for category, words in KEYWORD_DICTIONARY.items():
        html += f"<h3>{category}</h3>"
        html += f"<p>{', '.join(words)}</p>"
        
    html += "</body></html>"
    
    try:
        from weasyprint import HTML
        pdf_path = os.path.join(OUTPUT_DIR, "master_cv_database_v8.3.pdf")
        HTML(string=html).write_pdf(pdf_path)
        print(f"PDF saved to {pdf_path}")
    except ImportError:
        print("WeasyPrint not installed. Skipping PDF.")
    except Exception as e:
        print(f"PDF generation failed: {e}")

if __name__ == "__main__":
    generate_v83()
